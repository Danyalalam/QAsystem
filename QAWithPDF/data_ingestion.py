# data_ingestion.py

import logging
from llama_index.core import Document
from llama_index.core.node_parser import SimpleNodeParser
import io
from typing import List
from PyPDF2 import PdfReader
import docx

# Configure logging
logging.basicConfig(level=logging.INFO)

def load_documents_from_uploaded_files(uploaded_files) -> List[Document]:
    """
    Loads documents from uploaded files.

    Args:
        uploaded_files (List): List of uploaded files from Streamlit.

    Returns:
        List[Document]: A list of Document objects.
    """
    documents = []
    try:
        for uploaded_file in uploaded_files:
            file_type = uploaded_file.type
            file_name = uploaded_file.name
            logging.info(f"Processing file: {file_name} of type {file_type}")

            if file_type == 'text/plain':
                # Handle text files
                content = uploaded_file.read().decode('utf-8')
                documents.append(Document(text=content, doc_id=file_name))
            elif file_type == 'application/pdf':
                # Handle PDF files
                pdf_reader = PdfReader(uploaded_file)
                text = ''
                for page in pdf_reader.pages:
                    text += page.extract_text()
                documents.append(Document(text=text, doc_id=file_name))
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Handle Word documents (.docx)
                doc = docx.Document(uploaded_file)
                text = '\n'.join([para.text for para in doc.paragraphs])
                documents.append(Document(text=text, doc_id=file_name))
            else:
                logging.warning(f"Unsupported file type: {file_type}")
                continue
        return documents
    except Exception as e:
        logging.error(f"Error loading uploaded files: {e}")
        raise
